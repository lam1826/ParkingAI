from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
REPORT_PATH = ROOT / "HoSo_BaoCao_ParkingAI/04_BaoCao/BaoCao_DoAn_ParkingAI.docx"
TEST_REPORT_PATH = (
    ROOT
    / "HoSo_BaoCao_ParkingAI/05_GenAI_SDLC/05_GenAI_SoftwareDevelopment_functional-testing.docx"
)

COMPLETION_DATE = "23/08/2026"
BACKEND_RESULT = "79 passed, 1 warning trong 21,92 giây"
FRONTEND_RESULT = "4/4 kiểm thử tiện ích đạt"


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_cell_text(cell, text: str, *, bold: bool | None = None) -> None:
    paragraph = cell.paragraphs[0]
    set_paragraph_text(paragraph, text)
    for extra in cell.paragraphs[1:]:
        set_paragraph_text(extra, "")
    if bold is not None:
        for run in paragraph.runs:
            run.bold = bold


def replace_in_document(document, old: str, new: str) -> None:
    paragraphs = list(document.paragraphs)
    paragraphs.extend(
        paragraph
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )
    for paragraph in paragraphs:
        if old in paragraph.text:
            set_paragraph_text(paragraph, paragraph.text.replace(old, new))


def set_cell_margins(cell, top=55, start=70, bottom=55, end=70) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{margin}"))
        if element is None:
            element = OxmlElement(f"w:{margin}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_fixed_layout(table) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_table_widths(table, widths: list[float]) -> None:
    set_table_fixed_layout(table)
    for row in table.rows:
        for index, width in enumerate(widths):
            cell = row.cells[index]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def clear_paragraph_numbering(paragraph) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    numbering = paragraph_properties.find(qn("w:numPr"))
    if numbering is not None:
        paragraph_properties.remove(numbering)


def format_table(table, widths: list[float], font_size: float) -> None:
    set_table_widths(table, widths)
    repeat_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        prevent_row_split(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_borders(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                if row_index == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(font_size)
                    if row_index == 0:
                        run.bold = True


def save_atomic(document, path: Path) -> None:
    temporary_path = path.with_name(f".{path.stem}.tmp{path.suffix}")
    document.save(temporary_path)
    os.replace(temporary_path, path)


def finalize_main_report() -> None:
    document = Document(REPORT_PATH)

    cover = document.tables[0]
    set_cell_text(cover.rows[2].cells[1], COMPLETION_DATE)
    set_cell_text(cover.rows[4].cells[1], "Lý Trường Sơn; Lâm Duy Lập")
    set_cell_text(cover.rows[5].cells[1], "Chưa cung cấp trong hồ sơ dự án")

    for paragraph in document.paragraphs[12:23]:
        paragraph.style = document.styles["Normal"]
        clear_paragraph_numbering(paragraph)
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.space_after = Pt(3)

    for index, paragraph in enumerate(document.paragraphs[105:111], start=1):
        text = paragraph.text.strip()
        if not text.startswith(f"{index}."):
            set_paragraph_text(paragraph, f"{index}. {text}")
        paragraph.style = document.styles["Normal"]
        clear_paragraph_numbering(paragraph)
        paragraph.paragraph_format.left_indent = Inches(0.3)
        paragraph.paragraph_format.first_line_indent = Inches(-0.2)
        paragraph.paragraph_format.space_after = Pt(4)

    replacements = {
        "18/08/2026": COMPLETION_DATE,
        "79 kiểm thử backend đều đạt": "79 lượt kiểm thử backend và 4 kiểm thử tiện ích frontend đều đạt",
        "Tăng kiểm thử giao diện đầu-cuối và đo hiệu năng API.": "Mở rộng kiểm thử frontend sang component/E2E và đo hiệu năng API.",
        "Chạy pytest, ESLint và production build trước khi bàn giao.": "Chạy scripts/verify.ps1 để kiểm tra backend, frontend, ESLint và production build trước khi bàn giao.",
    }
    for old, new in replacements.items():
        replace_in_document(document, old, new)

    quality_table = document.tables[1]
    set_cell_text(
        quality_table.rows[5].cells[1],
        "pytest, Node test, ESLint, Vite build, GitHub Actions",
    )
    set_cell_text(
        quality_table.rows[5].cells[2],
        "Kiểm thử tự động, xác minh bản dựng và kiểm tra liên tục",
    )

    result_table = document.tables[11]
    result_rows = [
        ("Pytest backend", BACKEND_RESULT, "ĐẠT"),
        ("Node test frontend", FRONTEND_RESULT, "ĐẠT"),
        ("ESLint frontend", "Không có lỗi", "ĐẠT"),
        ("Vite production build", "2.106 module được xử lý", "ĐẠT"),
        ("Cảnh báo", "1 StarletteDeprecationWarning", "THEO DÕI"),
    ]
    while len(result_table.rows) < len(result_rows) + 1:
        result_table.add_row()
    for row, values in zip(result_table.rows[1:], result_rows):
        for cell, value in zip(row.cells, values):
            set_cell_text(cell, value)

    for table in document.tables:
        repeat_header(table.rows[0])
        for row in table.rows:
            prevent_row_split(row)

    save_atomic(document, REPORT_PATH)


ADDITIONAL_TEST_CASES = [
    (
        "TC-CHECKIN-09",
        "FR-05 (Xe vào/ra)",
        "Nhân viên chọn đích danh vị trí đỗ hợp lệ khi check-in.",
        "Vị trí đang hoạt động, còn trống và đúng loại xe.",
        'parking_slot_id hợp lệ; license_plate="43A-222.22".',
        "HTTP 201; trả đúng slot_id/slot_name; vị trí chuyển sang đã có xe.",
        "test_check_in.py::test_check_in_with_chosen_slot_success",
    ),
    (
        "TC-CHECKIN-10",
        "FR-05 (Xe vào/ra)",
        "Từ chối vị trí đỗ không tồn tại.",
        "User staff đã đăng nhập.",
        "parking_slot_id=999999.",
        'HTTP 404; detail chứa "không tồn tại".',
        "test_check_in.py::test_check_in_with_nonexistent_slot",
    ),
    (
        "TC-CHECKIN-11",
        "FR-05 (Xe vào/ra)",
        "Từ chối vị trí đỗ đã có xe.",
        "Vị trí hợp lệ nhưng is_occupied=True.",
        "parking_slot_id của vị trí đã có xe.",
        'HTTP 409; detail chứa "đã có xe".',
        "test_check_in.py::test_check_in_with_occupied_slot",
    ),
    (
        "TC-CHECKIN-12",
        "FR-05 (Xe vào/ra)",
        "Từ chối vị trí không hỗ trợ loại xe được chọn.",
        "Có vị trí xe máy; request chọn loại ô tô.",
        "vehicle_type_id không khớp loại của parking_slot_id.",
        "HTTP 400; không tạo phiên gửi xe.",
        "test_check_in.py::test_check_in_with_wrong_vehicle_type_slot",
    ),
    (
        "TC-CHECKOUT-07",
        "FR-05 (Xe vào/ra)",
        "Endpoint check-out phụ bỏ qua phí client và để server tính phí.",
        "Phiên active khoảng 30 phút; có bảng giá giờ.",
        "Client cố tình gửi parking_fee=0.",
        "HTTP 200; phí bằng giá một giờ; vị trí được giải phóng.",
        "test_check_out.py::test_crud_check_out_server_calculates_fee",
    ),
    (
        "TC-CHECKOUT-08",
        "FR-05 (Xe vào/ra)",
        "Từ chối check-out lần hai trên cùng một phiên.",
        "Phiên được check-out thành công ở lần đầu.",
        "Gọi lại endpoint check-out với cùng session_id.",
        "Lần hai trả HTTP 400; không tính phí hai lần.",
        "test_check_out.py::test_crud_check_out_already_completed",
    ),
    (
        "TC-FEE-07",
        "FR-05 (Tính phí)",
        "Tính phí đúng tại mốc chính xác một giờ.",
        "Có bảng giá HOURLY đang hiệu lực.",
        "Khoảng thời gian=3600 giây.",
        "Phí bằng giá một giờ.",
        "test_fee.py::test_calculate_fee_boundary_exact_hour",
    ),
    (
        "TC-FEE-08",
        "FR-05 (Tính phí)",
        "Làm tròn khoảng thời gian ngay dưới một giờ.",
        "Có bảng giá HOURLY đang hiệu lực.",
        "Khoảng thời gian=3599 giây.",
        "Phí bằng giá một giờ.",
        "test_fee.py::test_calculate_fee_boundary_just_under_hour",
    ),
    (
        "TC-FEE-09",
        "FR-05 (Tính phí)",
        "Làm tròn khoảng thời gian vừa vượt một giờ.",
        "Có bảng giá HOURLY đang hiệu lực.",
        "Khoảng thời gian=3601 giây.",
        "Phí bằng giá hai giờ.",
        "test_fee.py::test_calculate_fee_boundary_just_over_hour",
    ),
    (
        "TC-FEE-10",
        "FR-05 (Tính phí)",
        "Xử lý trường hợp vào và ra cùng thời điểm.",
        "Có bảng giá HOURLY đang hiệu lực.",
        "time_in=time_out.",
        "Phí bằng 0.",
        "test_fee.py::test_calculate_fee_zero_duration",
    ),
    (
        "TC-AI-13",
        "FR-08 (AI)",
        "Trả lỗi cấu hình rõ ràng khi thiếu GEMINI_API_KEY.",
        "Tạm thời đặt GEMINI_API_KEY rỗng.",
        "Gửi câu hỏi hợp lệ; gọi thống kê cơ bản sau đó.",
        "AI trả HTTP 503; thống kê cơ bản vẫn trả HTTP 200.",
        "test_ai.py::test_ai_missing_api_key_returns_503",
    ),
    (
        "TC-AI-14",
        "FR-08 (AI)",
        "Backend tự tổng hợp dữ liệu thật cho báo cáo ngày.",
        "Có một phiên gửi xe trong SQLite in-memory.",
        "Client chỉ gửi target_date, không gửi parking_stats.",
        "HTTP 200; prompt AI chứa total_vehicles_today.",
        "test_ai.py::test_ai_daily_report_server_side_aggregation",
    ),
    (
        "TC-AI-15",
        "FR-08 (AI)",
        "Backend tự tổng hợp dữ liệu cho gợi ý nhân sự.",
        "Có dữ liệu phiên gửi trong SQLite in-memory.",
        "Client gửi body rỗng.",
        "HTTP 200; prompt chứa hourly_traffic và occupancy_rate.",
        "test_ai.py::test_ai_staff_suggestion_server_side_aggregation",
    ),
]


def append_missing_test_rows(test_table, result_table) -> None:
    existing_ids = {row.cells[0].text.strip() for row in test_table.rows[1:]}
    for test_case in ADDITIONAL_TEST_CASES:
        if test_case[0] in existing_ids:
            continue
        row = test_table.add_row()
        for cell, value in zip(row.cells, test_case):
            set_cell_text(cell, value)

        result_row = result_table.add_row()
        result_values = (
            test_case[0],
            COMPLETION_DATE,
            "Kiểm thử tự động bằng pytest.",
            "Pass",
            "-",
            "Không có lỗi",
            "Đối chiếu mã nguồn.",
        )
        for cell, value in zip(result_row.cells, result_values):
            set_cell_text(cell, value)


def finalize_functional_test_report() -> None:
    document = Document(TEST_REPORT_PATH)

    for section in document.sections:
        if section.orientation != WD_ORIENT.LANDSCAPE:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.45)
        section.right_margin = Inches(0.45)

    hardware_text = (
        "Phần cứng kiểm thử: Intel Core i5-12450HX 64-bit, RAM 15,7 GB, ổ D: 244,1 GB "
        "(còn trống 86,3 GB); máy có kết nối mạng."
    )
    set_paragraph_text(document.paragraphs[8], hardware_text)

    hardware_table = document.tables[0]
    hardware_values = ("Intel Core i5-12450HX", "15,7 GB", "D: 244,1 GB", "64-bit")
    for cell, value in zip(hardware_table.rows[1].cells, hardware_values):
        set_cell_text(cell, value)

    software_table = document.tables[1]
    software_updates = {
        "SQLite (qua SQLAlchemy)": "SQLite 3.53.1; SQLAlchemy 2.0.51",
        "Visual Studio Code (hoặc IDE tương đương)": "Codex Desktop; PowerShell 7.6.4",
        "Hệ điều hành máy chạy kiểm thử": "Windows 11 Home Single Language 10.0.26200",
    }
    for row in software_table.rows[1:]:
        software_name = row.cells[0].text.strip()
        if software_name in software_updates:
            set_cell_text(row.cells[1], software_updates[software_name])

    test_table = document.tables[2]
    result_table = document.tables[3]
    append_missing_test_rows(test_table, result_table)

    additional_ids = {test_case[0] for test_case in ADDITIONAL_TEST_CASES}
    for row in result_table.rows[1:]:
        if row.cells[0].text.strip() in additional_ids:
            set_cell_text(row.cells[6], "Đối chiếu mã nguồn.")

    fee_row = next(row for row in test_table.rows if row.cells[0].text.strip() == "TC-FEE-03")
    set_cell_text(
        fee_row.cells[6],
        "test_fee.py::test_calculate_fee_monthly_pass — nhánh MonthlyPass thật được kiểm chứng độc lập tại TC-CHECKOUT-04.",
    )

    for row in result_table.rows[1:]:
        set_cell_text(row.cells[1], COMPLETION_DATE)
        if "không có người vận hành thủ công" in row.cells[2].text.lower():
            set_cell_text(row.cells[2], "Kiểm thử tự động bằng pytest.")

    summary = (
        "Ghi chú chung: Tài liệu truy vết 77 hàm test, tương ứng 79 lượt thực thi pytest; "
        "test_ai_reports_reject_empty_data được tham số hóa thành 3 trường hợp nên phát sinh thêm 2 lượt. "
        "Lệnh tái hiện từ thư mục gốc: .venv\\Scripts\\python.exe -m pytest -q. Kết quả ngày "
        f"{COMPLETION_DATE}: 79 passed, 1 warning in 21.92s. Cảnh báo StarletteDeprecationWarning "
        "không làm test thất bại. Toàn bộ test dùng SQLite in-memory (sqlite:///:memory:), tách biệt "
        "khỏi backend/database/parking.db. Frontend: 4/4 Node test đạt, ESLint không lỗi và Vite build "
        "thành công với 2.106 module."
    )
    summary_paragraph = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("Ghi chú chung:")
    )
    set_paragraph_text(summary_paragraph, summary)

    result_heading = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("3. Báo cáo kết quả test")
    )
    body = result_heading._p.getparent()
    previous = result_heading._p.getprevious()
    while previous is not None and previous.tag == qn("w:p") and not "".join(previous.itertext()).strip():
        candidate = previous.getprevious()
        body.remove(previous)
        previous = candidate
    result_heading.paragraph_format.page_break_before = True

    replace_in_document(document, "17/08/2026", COMPLETION_DATE)
    replace_in_document(document, "[CẦN BỔ SUNG]", "Đã xác minh ngày 23/08/2026")

    format_table(hardware_table, [2.7, 2.7, 2.7, 2.7], 9.5)
    format_table(software_table, [3.0, 4.1, 3.7], 9.0)
    format_table(test_table, [0.85, 1.35, 1.70, 1.45, 1.40, 2.35, 1.75], 7.5)
    format_table(result_table, [0.85, 1.00, 1.55, 0.65, 0.85, 2.70, 3.25], 8.0)

    save_atomic(document, TEST_REPORT_PATH)


def main() -> None:
    finalize_main_report()
    finalize_functional_test_report()
    print(f"Updated: {REPORT_PATH}")
    print(f"Updated: {TEST_REPORT_PATH}")


if __name__ == "__main__":
    main()
